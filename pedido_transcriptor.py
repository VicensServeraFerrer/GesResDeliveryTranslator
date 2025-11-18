import argparse
import pandas as pd
from urllib.parse import quote
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    Añade un hipervínculo clicable a un párrafo de python-docx.

    paragraph: objeto docx.paragraph.Paragraph
    url: enlace (string)
    text: texto visible del enlace
    color: color en hex (por defecto azul)
    underline: True/False para subrayado
    """
    # Relación de hipervínculo en el documento
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    # <w:hyperlink r:id="...">
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # <w:r>
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Color
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    rPr.append(color_el)

    # Subrayado
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    new_run.append(rPr)

    # Texto del enlace
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)

    # Lo enganchamos todo
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return paragraph


def generate_messages(
    input_excel,
    output_file,
    sheet_name="Pedido",
    header_row=6,
    header_text="Hola, os paso el pedido del Celler de Randa para esta semana:",
):
    # Leer Excel
    df = pd.read_excel(input_excel, sheet_name=sheet_name, header=header_row)

    # Filtrar filas válidas
    df = df.dropna(subset=["Proveedor (auto)", "Producto (auto)"])

    # Opcional: filtrar cantidades 0 o NaN
    if "Cantidad" in df.columns:
        df = df[df["Cantidad"].fillna(0) != 0]

    # Agrupar por proveedor
    grupos = df.groupby("Proveedor (auto)")

    # Crear documento Word
    doc = Document()

    for proveedor, datos in grupos:
        lineas = []
        telefono = ""

        for _, fila in datos.iterrows():
            producto = str(fila["Producto (auto)"])
            cantidad = fila.get("Cantidad", "")
            unidad = fila.get("Unidad (auto)", "")

            # Cogemos el teléfono de la primera fila que tenga algo
            if not telefono:
                telefono = str(fila.get("Telefono", "") or "")

            linea = f"- {producto}: {cantidad} {unidad}".strip()
            lineas.append(linea)

        cuerpo = "\n".join(lineas)

        # Cabecera configurable
        cabecera = header_text.strip()
        if cabecera and not cabecera.endswith(":"):
            cabecera += ":"

        mensaje_completo = f"{cabecera}\n\n{cuerpo}"

        # Título con el nombre del proveedor
        doc.add_heading(str(proveedor), level=2)
        doc.add_paragraph(mensaje_completo)

        # Link de WhatsApp como hipervínculo (si hay teléfono)
        if telefono:
            encoded = quote(mensaje_completo)
            wa_link = f"https://wa.me/{telefono}?text={encoded}"

            p = doc.add_paragraph()
            p.add_run("WhatsApp: ")
            add_hyperlink(p, wa_link, "abrir conversación")

        # Salto de página entre proveedores
        doc.add_paragraph()

    # Guardar en la ruta que nos pasa Node
    doc.save(output_file)


def main():
    parser = argparse.ArgumentParser(description="Generar DOCX con mensajes de pedidos.")
    parser.add_argument("--input", required=True, help="Ruta del Excel de entrada")
    parser.add_argument("--output", required=True, help="Ruta del DOCX de salida")
    parser.add_argument("--sheet", default="Pedido", help="Nombre de la hoja (por defecto: Pedido)")
    parser.add_argument(
        "--header-row",
        type=int,
        default=6,
        help="Fila de cabecera (0-index, por defecto 6)",
    )
    parser.add_argument(
        "--header-text",
        type=str,
        default="Hola, os paso el pedido del Celler de Randa para esta semana:",
        help="Texto de cabecera del mensaje",
    )

    args = parser.parse_args()

    generate_messages(
        input_excel=args.input,
        output_file=args.output,
        sheet_name=args.sheet,
        header_row=args.header_row,
        header_text=args.header_text,
    )


if __name__ == "__main__":
    main()

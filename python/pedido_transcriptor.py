import argparse
import pandas as pd
from utils.add_hyperlink import add_hyperlink
from urllib.parse import quote
from docx import Document

def generate_messages(
    input_excel,
    output_file,
    sheet_name="Pedido",
    header_row=7,
    header_text="Hola, os paso el pedido del Celler de Randa para esta semana:",
):
    df = pd.read_excel(input_excel,sheet_name=sheet_name, header=None)

    # Search header and cut dataframe
    start_row = df[df.eq("Proveedor (auto)").any(axis=1)].index[0]
    df = df.iloc[start_row:].reset_index(drop=True)

    # Use cell as header
    df.columns = df.iloc[0]
    df = df.iloc[1:].reset_index(drop=True)

    # Filtrate empty rows
    df = df.dropna(subset=["Proveedor (auto)", "Producto (auto)"])

    #Filtrate Cantidad 
    df["Cantidad"] = (pd.to_numeric(df["Cantidad"], errors="coerce"))
    df = df[df["Cantidad"].fillna(0) != 0]

    grupos = df.groupby("Proveedor (auto)")

    doc = Document()

    for proveedor, datos in grupos:
        lineas = []
        telefono = ""

        for _, fila in datos.iterrows():
            producto = str(fila["Producto (auto)"])
            cantidad = int(fila.get("Cantidad", ""))
            unidad = fila.get("Unidad (auto)", "")

            # Cogemos el teléfono de la primera fila que tenga algo
            if not telefono:
                telefono = str(fila.get("Telefono", "") or "")

            linea = f"- {producto}: {cantidad} {unidad}".strip()
            lineas.append(linea)

        cuerpo = "\n".join(lineas)

        # Cabecera configurable
        cabecera = header_text.strip()
        if cabecera and not cabecera.endswith(":"):
            cabecera += ":"

        mensaje_completo = f"{cabecera}\n\n{cuerpo}"

        # Título con el nombre del proveedor
        doc.add_heading(str(proveedor), level=2)
        doc.add_paragraph(mensaje_completo)

        # Link de WhatsApp como hipervínculo (si hay teléfono)
        if telefono:
            encoded = quote(mensaje_completo)
            wa_link = f"https://wa.me/{telefono}?text={encoded}"

            p = doc.add_paragraph()
            p.add_run("WhatsApp: ")
            add_hyperlink(p, wa_link, "abrir conversación")

        # Salto de página entre proveedores
        doc.add_paragraph()

    # Guardar en la ruta que nos pasa Node
    doc.save(output_file)


def main():
    parser = argparse.ArgumentParser(description="Generar DOCX con mensajes de pedidos.")
    parser.add_argument("--input", required=True, help="Ruta del Excel de entrada")
    parser.add_argument("--output", required=True, help="Ruta del DOCX de salida")
    parser.add_argument("--sheet", default="Pedido", help="Nombre de la hoja (por defecto: Pedido)")
    parser.add_argument(
        "--header-row",
        type=int,
        default=7,
        help="Fila de cabecera (0-index, por defecto 6)",
    )
    parser.add_argument(
        "--header-text",
        type=str,
        default="Hola, os paso el pedido para esta semana:",
        help="Texto de cabecera del mensaje",
    )

    args = parser.parse_args()

    generate_messages(
        input_excel=args.input,
        output_file=args.output,
        sheet_name=args.sheet,
        header_row=args.header_row,
        header_text=args.header_text,
    )


if __name__ == "__main__":
    main()
